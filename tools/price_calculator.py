from docx import Document

def return_product(product):
    print('-=-=-=-')
    print('Product Name:', product['name'])
    print('-=-=-=-')
    print(product['description'])
    print('-=-=-=-')
    print('Supplies:')
    print('-=-=-=-')
    total_price = 0
    for supply in product['supplies']:
        print('Name:', supply['name'])
        print('Description:', supply['description'])
        print('Quantity:', supply['quantity'])
        print('Price:', supply['price'])
        print('Cost per Unit:', supply['cost_per_unit'])
        total_price = total_price + supply['price']
        print('-=-=-=-')
    print('Instructions:')
    for i, instruction in enumerate(product['instructions'], start=1):
        print(f'{i}.) {instruction}')
    print('-=-=-=-')
    print('Total cost: $', str(total_price))
    print('-=-=-=-')
    cost_per_product = (total_price / product['quantity'])
    print('Makes ' + str(product['quantity']) + 'of ' + product['name'] + ' for $' + str(cost_per_product))
    print('-=-=-=-')
    print('Weight of individual product: ')
    print('-=-=-=-')

    doc = Document()

    doc.add_heading(product['name'].title(), level=1)
    doc.add_heading('Supplies:', level=4)

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Quantity'
    hdr_cells[3].text = 'Price'
    hdr_cells[4].text = 'Cost per Unit'

    for supply in product['supplies']:
        row_cells = table.add_row().cells
        row_cells[0].text = supply['name']
        row_cells[1].text = supply['description']
        row_cells[2].text = (str(supply['quantity']) + supply['unit_of_measure'])
        row_cells[3].text = str(supply['price'])
        row_cells[4].text = (str(supply['cost_per_unit']))

    doc.add_heading('Total cost to make ' + str(product['quantity']) + ' ' + product['name'] + ' that weigh ' + product['net_weight'] + ' each: $' + str(total_price), level=2)
    doc.add_heading('Cost per product: $' + str(cost_per_product), level=3)

    while True:
        upcharge = (int(input('Upcharge percentage: ')))
        new_price = cost_per_product * ((upcharge / 100) + 1)
        choice = input('Do you want to upcharge ' + str(upcharge) + '% ? It will bring the price from $' + str(cost_per_product) + ' to $' + str(new_price) + '. Y/N: ')

        if choice.lower() == 'y':
            doc.add_heading('Total price for product after ' + str(upcharge) + '% ' + 'upcharge: $' + str(new_price), level=3)
            break
        else:
            print('OK.')

    doc.add_heading('Description: ', level=2)
    doc.add_paragraph(product['description'])

    doc.add_heading('Instructions: ', level=2)
    for i, instruction in enumerate(product['instructions'], start=1):
        doc.add_heading(f'Step {i}:', level=3)
        doc.add_paragraph(f'{instruction}')

    doc.save(product['name'] + '_PRODUCT.docx')
    print('Saved as ' + product['name'] + '_PRODUCT.docx')


def add_product():
    product = {}
    product['name'] = input('Name of product: ').title()
    product['description'] = input('Description: ')
    product['quantity'] = int(input('How many products it makes: '))

    unit = input('Unit of measurement: ').lower()
    net_weight = float(input('Weight in ' + unit + ' of individual product: '))

    product['net_weight'] = (str(net_weight) + unit)
    product['supplies'] = []
    product['instructions'] = []

    choice = input('Add supply? Y/N: ').lower()

    if choice == 'y':
        while True:
            add_supply(product['supplies'])
            add_more = input('Add another supply? Y/N: ').lower()
            if add_more != 'y':
                add_instructions(product, product['instructions'])
                break
    else:
        print('OK.')

def add_instructions(product, instructions):
    step = 0

    while True:
        step = step + 1
        instructions_info = input('Step ' + str(step) + '.) ')
        instructions.append(instructions_info)

        choice = input('Add another? Y/N: ').lower()
        if choice != 'y':
            return_product(product)
            break


def add_supply(supplies):
    supply_info = {}
    supply_info['name'] = input('Name of supply: ').title()
    supply_info['description'] = input('Description: ').capitalize()

    print('1.) Gram')
    print('2.) Ounce')
    print('3.) Milliliter')
    print('4.) Pieces')

    choice = int(input('Pick one: '))

    uom = ''

    if choice in [1, 2, 3, 4]:
        if choice == 1:
            uom = 'g'
        elif choice == 2:
            uom = 'oz'
        elif choice == 3:
            uom = 'ml'
        elif choice == 4:
            uom = 'pc'
    else:
        print('Invalid input.')

    qty = float(input('Amount of ' + supply_info['name'] + ' in ' + uom))

    price = float(input('Cost for ' + str(qty) + uom + ' of ' + supply_info['name']))

    cost_per_unit = (price / qty)

    print('Cost per unit for ' + supply_info['name'] + ': $' + str(cost_per_unit))

    supply_info['quantity'] = qty
    supply_info['unit_of_measure'] = uom
    supply_info['price'] = price
    supply_info['cost_per_unit'] = cost_per_unit

    supplies.append(supply_info)

choice = input('Add product? Y/N: ').lower()

if choice == 'y':
    add_product()
